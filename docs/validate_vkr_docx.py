#!/usr/bin/env python3
"""Validate common Russian VKR/DOCX formatting rules.

Usage:
    python3 validate_vkr_docx.py path/to/report.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def _has_prior_ref(kind: str, number: int, text: str) -> bool:
    stem = "таблиц" if kind == "table" else "рисунк"
    if re.search(rf"{stem}\w*\s+{number}(?!\d)", text, re.IGNORECASE):
        return True
    for match in re.finditer(rf"{stem}\w*\s+(\d+)\s*[–-]\s*(\d+)", text, re.IGNORECASE):
        start, end = int(match.group(1)), int(match.group(2))
        if start <= number <= end:
            return True
    return False


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: validate_vkr_docx.py path/to/report.docx", file=sys.stderr)
        return 2

    path = Path(sys.argv[1])
    doc = Document(path)
    errors: list[str] = []
    paragraphs = list(doc.paragraphs)
    texts = [p.text.strip() for p in paragraphs]

    table_captions: list[tuple[int, int, str]] = []
    figure_captions: list[tuple[int, int, str]] = []

    for index, text in enumerate(texts):
        table_match = re.match(r"^Таблица\s+(\d+)\s+[–-]\s+(.+)", text)
        if table_match:
            table_captions.append((index, int(table_match.group(1)), text))

        figure_match = re.match(r"^Рисунок\s+(\d+)\s+[–-]\s+(.+)", text)
        if figure_match:
            figure_captions.append((index, int(figure_match.group(1)), text))

    for _, _, caption in table_captions + figure_captions:
        if " - " in caption:
            errors.append(f"short dash in caption: {caption}")

    for index, number, caption in table_captions:
        prior_text = "\n".join(texts[:index])
        if not _has_prior_ref("table", number, prior_text):
            errors.append(f"no prior table reference: {caption}")
        if paragraphs[index].paragraph_format.keep_with_next is not True:
            errors.append(f"table caption is not keep_with_next: {caption}")

    for index, number, caption in figure_captions:
        prior_text = "\n".join(texts[:index])
        if not _has_prior_ref("figure", number, prior_text):
            errors.append(f"no prior figure reference: {caption}")
        if index == 0:
            errors.append(f"figure caption has no previous image paragraph: {caption}")
            continue
        previous = paragraphs[index - 1]
        if previous.paragraph_format.keep_with_next is not True:
            errors.append(f"image paragraph is not keep_with_next before: {caption}")
        space_after = previous.paragraph_format.space_after
        if space_after is None or abs(space_after.pt - 10) > 0.2:
            actual = "none" if space_after is None else f"{space_after.pt:.1f}"
            errors.append(f"image paragraph space_after is not 10 pt before {caption}: {actual}")

    for table_index, table in enumerate(doc.tables, start=1):
        header = table.rows[0]
        header_props = header._tr.trPr
        if header_props is None or header_props.find(qn("w:tblHeader")) is None:
            errors.append(f"table {table_index}: header row repeat is not enabled")

        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    if (
                        paragraph_format.first_line_indent is not None
                        and abs(paragraph_format.first_line_indent.pt) > 0.2
                    ):
                        errors.append(f"table {table_index}: paragraph first-line indent in row {row_index + 1}")
                    if (
                        paragraph_format.left_indent is not None
                        and abs(paragraph_format.left_indent.pt) > 0.2
                    ):
                        errors.append(f"table {table_index}: paragraph left indent in row {row_index + 1}")
                    if paragraph_format.line_spacing is not None and paragraph_format.line_spacing != 1:
                        errors.append(f"table {table_index}: line spacing is not single in row {row_index + 1}")
                    for run in paragraph.runs:
                        if not run.text.strip():
                            continue
                        if run.font.size is not None and abs(run.font.size.pt - 12) > 0.2:
                            errors.append(f"table {table_index}: font size is not 12 pt")
                        if row_index == 0 and run.bold is not True:
                            errors.append(f"table {table_index}: header text is not bold")

    print(f"tables: {len(doc.tables)}, table captions: {len(table_captions)}")
    print(f"figures: {len(doc.inline_shapes)}, figure captions: {len(figure_captions)}")

    if errors:
        print(f"errors: {len(errors)}")
        for error in errors:
            print(f"ERROR: {error}")
        return 1

    print("errors: 0")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
